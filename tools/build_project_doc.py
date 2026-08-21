from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "ROV_Digital_Twin_Project_Outline.docx"
METRICS = json.loads((ROOT / "artifacts" / "training_metrics.json").read_text(encoding="utf-8"))

NAVY = "17324D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8F1F8"
PALE = "F3F6F9"
MUTED = "5B6573"
WHITE = "FFFFFF"
RED = "A63D40"
GREEN = "1F7A4D"
BLACK = "1E1E1E"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_font(run, name="Calibri", size=None, color=BLACK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 9, NAVY),
        ("Heading 2", 13, 13, 6, BLUE),
        ("Heading 3", 11.5, 9, 4, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18


def add_header_footer(section) -> None:
    for header in (section.header, section.even_page_header, section.first_page_header):
        hp = header.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = hp.add_run("ROV DIGITAL TWIN  /  INTELLIGENCE STACK")
        set_font(r, size=8.5, color=MUTED, bold=True)
        p_pr = hp._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), "CFD8E3")
        border.append(bottom)
        p_pr.append(border)

    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = fp.add_run("Teknik Proje Outline  |  Ağustos 2026  |  Sayfa ")
        set_font(r, size=8.5, color=MUTED)
        add_page_number(fp)


def para(doc, text, *, bold_prefix=None, italic=False, align=None, before=0, after=6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_font(run, italic=italic)


def bullet(doc, text, level=0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.65)
    set_font(p.add_run(text))


def number(doc, text) -> None:
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run(text))


def callout(doc, label: str, text: str, fill=LIGHT_BLUE, accent=BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(label.upper() + "  ")
    set_font(r, size=9, color=accent, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5, color=BLACK)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=9.5, color=WHITE, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[idx], PALE)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), size=9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def page_break(doc) -> None:
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    configure_styles(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("TEKNİK PROJE OUTLINE"), size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    set_font(p.add_run("Su Altı Robotu Dijital İkizi"), size=28, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_font(p.add_run("RL kontrolü, weak-point sınıflandırması, uzman LLM ve emniyet kapılı karar ajanı"), size=14, color=MUTED)
    callout(doc, "Teslimat kapsamı", "Çalışan Python referans hattı, örnek veri/model/karar çıktıları, Unity ve ROS 2 entegrasyon iskeleti, LLM instruction veri seti ve üretime geçiş planı.")
    para(doc, "Sürüm: 0.1.0", align=WD_ALIGN_PARAGRAPH.CENTER, before=32, after=3)
    para(doc, "Tarih: 21 Ağustos 2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    para(doc, "Durum: Reference implementation / MVP", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    page_break(doc)

    doc.add_heading("1. Yönetici özeti", level=1)
    para(doc, "Bu proje, Unity içinde çalışan 6-DOF su altı araç simülasyonunu; ROS 2 telemetrisi, RL tabanlı duty kontrolü, arıza/zayıf-nokta tanılama ve emniyet kontrollü bir karar katmanıyla birleştirir. İlk sürümün amacı gerçek aracı otonom yönetmek değil; ölçülebilir, tekrarlanabilir ve sahaya taşınabilir bir dijital ikiz zekâ omurgası oluşturmaktır.")
    callout(doc, "Ana karar", "LLM yalnızca açıklama ve kontrol listesi üretir. Nihai komut uygunluğu; deterministik eşikler, model güveni ve operatör onayıyla yönetilen SafetyDecisionAgent tarafından belirlenir.", fill="FFF5E8", accent=RED)
    doc.add_heading("1.1 Başarı tanımı", level=2)
    bullet(doc, "Aynı şema ile sentetik, SIL/HIL ve saha telemetrisini işleyebilmek.")
    bullet(doc, "Nominal durum ile thruster degradation, buoyancy imbalance, sensor drift ve hydrodynamic drag sınıflarını ayırmak.")
    bullet(doc, "Her tahminde sınıf, güven, kanıt, risk, eylem ve otonom yürütme izni üretmek.")
    bullet(doc, "RL politikasını duty hedefleri ve enerji/emniyet cezalarıyla eğitip ONNX olarak Unity'ye taşımak.")
    bullet(doc, "Karar ve model sürümlerini denetlenebilir biçimde kaydetmek.")

    doc.add_heading("2. Sistem sınırları ve ilkeler", level=1)
    add_table(doc, ["Kapsamda", "Kapsam dışında / sonraki faz"], [
        ["Unity Rigidbody tabanlı hidrodinamik referans model", "CFD düzeyinde akış çözümü"],
        ["Tekli weak-point sınıflandırması", "Aynı anda çoklu arıza teşhisi"],
        ["LoRA instruction verisi ve eğitim giriş noktası", "Sertifikalı otonom komut verme"],
        ["ROS 2 JSON telemetri adaptörü", "Üretim mesaj tipleri ve gerçek araç sürücüsü"],
        ["Sentetik benchmark ve testler", "Saha performansı iddiası"],
    ], [4680, 4680])
    para(doc, "Önemli sınırlama: sentetik benchmark üzerindeki yüksek metrikler yalnızca yazılım hattının çalıştığını gösterir. Saha doğruluğu, mission/time bazlı ayrılmış gerçek veri ve kör HIL test setiyle yeniden ölçülmelidir.", bold_prefix="Önemli sınırlama:")

    page_break(doc)
    doc.add_heading("3. Uçtan uca mimari", level=1)
    para(doc, "Veri akışı aşağıdaki sahiplik sırasını izler:")
    for step in (
        "Unity/gerçek ROV sensörleri telemetri üretir; ROS 2 zaman damgası ve görev bağlamını taşır.",
        "Şema katmanı birimleri doğrular ve 14 sayısal özelliği ortak vektöre dönüştürür.",
        "Weak-point modeli olasılık dağılımı ve en olası sınıfı üretir.",
        "Alan danışmanı telemetri kanıtlarını ve kontrol listesini oluşturur; üretimde bu arayüz RAG/fine-tuned LLM ile değiştirilebilir.",
        "SafetyDecisionAgent kritik eşikleri ve güven kapısını uygular.",
        "ROS gateway yalnızca izinli niyetleri operatör veya araç kontrol katmanına iletir.",
    ):
        number(doc, step)
    add_table(doc, ["Katman", "Ana bileşen", "Sorumluluk", "Fail-safe"], [
        ["Fizik", "Unity 6-DOF", "Kaldırma, drag, thruster kuvveti", "Simülasyon sınırları"],
        ["Kontrol", "ML-Agents PPO", "Duty politikası", "Episode sonlandırma"],
        ["Tanı", "Softmax model", "Weak-point olasılıkları", "Güven eşiği"],
        ["Açıklama", "DomainAdvisor / LLM", "Kanıt ve bakım kontrolü", "Komut yetkisi yok"],
        ["Karar", "SafetyDecisionAgent", "Risk ve eylem niyeti", "Hold / abort / review"],
        ["Entegrasyon", "ROS 2", "Telemetri ve karar mesajları", "Gateway izin listesi"],
    ], [1400, 2050, 3430, 2480])

    doc.add_heading("4. Veri stratejisi", level=1)
    doc.add_heading("4.1 Kaynaklar", level=2)
    bullet(doc, "Sentetik: fault injection ile dengeli başlangıç verisi; CI ve regresyon için.")
    bullet(doc, "SIL: Unity sahneleri, farklı akıntı, yük, görev ve sensör gürültüsü senaryoları.")
    bullet(doc, "HIL: gerçek kontrolcü, ESC ve sensörlerin kapalı çevrim kayıtları.")
    bullet(doc, "Saha: görev/araç/sensör kalibrasyonu ve operatör etiketleriyle versioned kayıt.")
    page_break(doc)
    doc.add_heading("4.2 Veri sözlüğü", level=2)
    add_table(doc, ["Grup", "Alanlar", "Birim / aralık", "Tanısal amaç"], [
        ["Görev", "timestamp, mission_id, duty", "s / kimlik / enum", "Bağlam ve doğru split"],
        ["Hareket", "depth, depth_error, speed, vertical_speed", "m, m/s", "Konum ve performans"],
        ["Tutum", "roll, pitch, yaw_rate", "deg, deg/s", "Dengesizlik/asimetri"],
        ["Güç", "current, voltage, temperature", "A, V, °C", "Yük ve sağlık"],
        ["Thruster", "command_mean, response_ratio", "0–1 / oran", "Aktüatör verimi"],
        ["Sensör", "IMU-depth disagreement, DVL quality", "m / 0–1", "Drift ve güven"],
    ], [1250, 2910, 1920, 3280])
    doc.add_heading("4.3 Etiketleme ve kalite", level=2)
    bullet(doc, "Etiket kaynağı: injected fault ID, bakım kaydı, test protokolü ve operatör doğrulaması ayrı alanlarda tutulmalı.")
    bullet(doc, "Eksik veri, zaman kayması, doygunluk ve dropout maskeleri model girdisine eklenmeli.")
    bullet(doc, "Train/validation/test ayrımı satır bazlı değil mission_id ve zaman bloğu bazlı yapılmalı.")
    bullet(doc, "Veri kartı; araç konfigürasyonu, su koşulları, yazılım sürümü ve bilinen boşlukları içermeli.")

    doc.add_heading("5. Weak-point modeli ve eğitim", level=1)
    para(doc, "Referans sınıflandırıcı; standartlaştırılmış ham özelliklere mutlak-değer özelliklerini ekleyen, L2 düzenlemeli çok sınıflı softmax modelidir. Bu seçim CI ortamında harici ML bağımlılığı olmadan eğitilebilirlik ve model JSON'unun kolay denetlenmesi için yapılmıştır.")
    add_table(doc, ["Sınıf", "Başlıca sinyal", "Önerilen güvenlik tepkisi"], [
        ["nominal", "Düşük hata, sağlıklı sensör/response", "Göreve devam + trend izleme"],
        ["thruster_degradation", "Düşük response, yüksek akım, yaw", "Degraded mode + operatör"],
        ["buoyancy_imbalance", "Dikey hız/depth error, pitch/roll", "Depth hold + trim kontrolü"],
        ["sensor_drift", "Sensör anlaşmazlığı, düşük DVL quality", "Cross-check + kaynağı reddet"],
        ["hydrodynamic_drag", "Yüksek akım, düşük hız, orta response", "Hızı azalt + tether/fouling kontrolü"],
    ], [2150, 3430, 3780])
    doc.add_heading("5.1 Eğitim ve değerlendirme", level=2)
    bullet(doc, "Dengeli 1.000 satırlık örnek veri; seed=42.")
    bullet(doc, "Stratified %80 train / %20 test; 180 epoch.")
    bullet(doc, f"Referans sonuç: accuracy={METRICS['accuracy']:.3f}, macro-F1={METRICS['macro_f1']:.3f}.")
    bullet(doc, "Sonuç sentetik sınıfların kontrollü ayrışmasını ölçer; saha genellemesi olarak yorumlanmamalıdır.")
    doc.add_heading("5.2 Üretim metrikleri", level=2)
    bullet(doc, "Sınıf bazlı recall: özellikle sensor drift ve thruster degradation için emniyet hedefi.")
    bullet(doc, "Expected Calibration Error ve reliability curve: güven kapısını doğrulamak için.")
    bullet(doc, "False-negative başına görev etkisi ve operatör override oranı.")
    bullet(doc, "Araç, görev, derinlik, akıntı ve sensör sürümü segmentlerinde drift.")

    doc.add_heading("6. RL duty kontrolü", level=1)
    add_table(doc, ["Duty", "Observation", "Reward", "Termination"], [
        ["Station keeping", "Hedef ofseti, hız, açısal hız", "Pozisyon + enerji", "Sapma/timeout"],
        ["Pipeline tracking", "Hat göreli poz/heading, DVL", "Takip hatası + pürüzsüzlük", "Hat kaybı/çarpışma"],
        ["Target waypoint", "Hedef vektörü, pose, akıntı", "İlerleme + enerji", "Hedef/timeout"],
    ], [1970, 2870, 2440, 2080])
    para(doc, "Eğitim ilerlemesi önce kolay koşullarda başlar; akıntı, sensör gürültüsü, drag, payload ve thruster verimi curriculum/domain randomization ile kademeli artırılır. Politika ONNX'e çevrilmeden önce holdout sahneler ve adversarial fault injection ile değerlendirilir.")

    doc.add_heading("7. Uzmanlaşmış LLM", level=1)
    para(doc, "LLM'nin görevi sınıflandırıcının yerine geçmek değil; telemetriyi bakım/operasyon bilgisiyle ilişkilendirerek kısa gerekçe, kontrol sırası ve operatör için açıklama üretmektir.")
    doc.add_heading("7.1 Veri ve yöntem", level=2)
    bullet(doc, "Instruction formatı: system safety policy + telemetry JSON + tanı/öneri yanıtı.")
    bullet(doc, "Başlangıç modeli: küçük instruct model; LoRA ile alan uyarlaması.")
    bullet(doc, "Bilgi tabanı: araç el kitabı, duty prosedürleri, bakım kayıtları, sensör datasheet'leri ve olay raporları.")
    bullet(doc, "RAG zorunluluğu: araç sürümüne özgü limitler ve bakım adımları kaynaklı verilmelidir.")
    doc.add_heading("7.2 Guardrail ve değerlendirme", level=2)
    add_table(doc, ["Kontrol", "Kabul kriteri", "Başarısızlık tepkisi"], [
        ["Komut yetkisi", "LLM raw actuator output üretmez", "Çıktıyı blokla"],
        ["Kanıt", "Her öneri en az bir telemetry sinyali içerir", "Operatör review"],
        ["Kaynak", "Araç-spesifik iddia retrieval kaynağı taşır", "Belirsizliği belirt"],
        ["Tutarlılık", "Aynı girdi güvenlik sınıfını değiştirmez", "Deterministik ajanı esas al"],
        ["Red-team", "Prompt injection komut katmanına ulaşmaz", "Gateway izolasyonu"],
    ], [1850, 3950, 3560])

    doc.add_heading("8. Karar ajanı", level=1)
    para(doc, "SafetyDecisionAgent üç sinyali birleştirir: model olasılıkları, kritik telemetri eşikleri ve alan danışmanı kanıtları. Çıktı şeması weak_point, confidence, risk_level, action, autonomous_execution_allowed, evidence, recommended_checks ve probabilities alanlarını taşır.")
    add_table(doc, ["Koşul", "Risk", "Eylem", "Otonom yürütme"], [
        ["Nominal + güven ≥ 0,72 + kritik limit yok", "low", "continue_mission", "Evet"],
        ["Non-nominal + yeterli güven", "high", "Degraded/hold/reduce + review", "Hayır"],
        ["Güven < 0,72", "uncertain", "hold_position + review", "Hayır"],
        ["Depth/pitch/voltage/temperature kritik", "critical", "abort_and_surface", "Hayır"],
    ], [2740, 1420, 3220, 1980])
    callout(doc, "Fail-safe", "Ağ, LLM veya sınıflandırıcı kullanılamıyorsa araç son bilinen komuta devam etmez; yerel kontrolcü tanımlı safe state'e geçer ve operatör uyarılır.", fill="FDEEEE", accent=RED)

    doc.add_heading("9. Unity ve ROS 2 entegrasyonu", level=1)
    doc.add_heading("9.1 Unity", level=2)
    bullet(doc, "Hydrodynamics6Dof: kaldırma, eksen bazlı karesel lineer ve açısal drag.")
    bullet(doc, "DutyDefinition: StationKeeping, PipelineTracking ve TargetWaypoint parametreleri.")
    bullet(doc, "ROVRLAgent: observation/action/reward yaşam döngüsü ve thruster kuvvetleri.")
    bullet(doc, "Sonraki adım: URDF importer, added-mass matrisi, thruster curve ve sensör noise profilinin kalibrasyonu.")
    doc.add_heading("9.2 ROS 2", level=2)
    bullet(doc, "Girdi topic: `/rov/telemetry_json`; çıktı: `/rov/diagnostic_decision`.")
    bullet(doc, "Referans node std_msgs/String kullanır; üretimde timestamp'li özel mesaj ve QoS profili gerekir.")
    bullet(doc, "Komut topic'i tanı node'undan ayrılmalı; gateway yalnızca imzalı/izinli eylem niyetlerini kabul etmelidir.")
    doc.add_heading("9.3 Sim-to-real", level=2)
    bullet(doc, "Fizik parametrelerini havuz/deniz deneyiyle system identification üzerinden ayarla.")
    bullet(doc, "Sensor noise, latency ve dropout dağılımlarını saha loglarından öğren.")
    bullet(doc, "SIL → HIL → tethered wet test → sınırlı görev sırasını uygula.")
    bullet(doc, "Shadow mode'da kararları operatör davranışıyla karşılaştır; komut yetkisini en son aç.")

    page_break(doc)
    doc.add_heading("10. MLOps, izlenebilirlik ve güvenlik", level=1)
    bullet(doc, "Her model artefaktı: veri sürümü, feature schema, seed, commit SHA, metrikler ve onay kaydı.")
    bullet(doc, "Runtime log: input window hash, model version, tahmin dağılımı, eşik sonucu, karar ve override.")
    bullet(doc, "Canary/shadow rollout; performans veya kalibrasyon bozulursa otomatik rollback.")
    bullet(doc, "ROS ağı segmentasyonu, least privilege, signed deployment artefaktı ve secret yönetimi.")
    bullet(doc, "Kişisel veri beklenmese de görev koordinatları ve operasyon logları erişim kontrollü tutulmalı.")

    doc.add_heading("11. Yol haritası ve iş paketleri", level=1)
    add_table(doc, ["Faz", "Süre", "Çıktı", "Çıkış kapısı"], [
        ["0 — Temel MVP", "1–2 hf", "Repo, sentetik veri, model, ajan, Unity/ROS iskeleti", "CI yeşil; demo tekrar üretilebilir"],
        ["1 — Kalibrasyon", "2–4 hf", "SIL senaryoları ve sensor/thruster profilleri", "Referans manevralar hata bütçesinde"],
        ["2 — Gerçek veri", "3–6 hf", "HIL/saha dataset v1 ve veri kartı", "Mission-split benchmark onaylı"],
        ["3 — RL doğrulama", "4–8 hf", "Duty policy, curriculum, ONNX", "Holdout görev başarı ve enerji hedefi"],
        ["4 — LLM/RAG", "3–5 hf", "Kaynaklı uzman danışman", "Safety eval ve red-team geçer"],
        ["5 — Wet test", "2–6 hf", "Shadow/degraded pilot", "Operatör ve emniyet kurulu onayı"],
    ], [1260, 1100, 3820, 3180])
    doc.add_heading("11.1 Rol önerisi", level=2)
    bullet(doc, "Robotics/Unity: fizik, URDF, sensör ve ML-Agents sahneleri.")
    bullet(doc, "ML engineer: dataset, classifier, RL training, değerlendirme ve deployment.")
    bullet(doc, "ROV/controls engineer: limitler, system identification, HIL ve fail-safe.")
    bullet(doc, "LLM engineer: instruction/RAG, eval setleri, guardrail ve observability.")
    bullet(doc, "Operator/safety owner: duty kabul kriterleri, test yetkisi ve go/no-go.")

    page_break(doc)
    doc.add_heading("12. Risk kaydı", level=1)
    add_table(doc, ["Risk", "Olasılık / Etki", "Azaltım", "Sahip"], [
        ["Sim-real gap", "Yüksek / Yüksek", "Domain randomization + HIL + shadow", "Controls/ML"],
        ["Yanlış negatif arıza", "Orta / Kritik", "Recall hedefi + conservative gate", "Safety/ML"],
        ["Sensör zaman kayması", "Yüksek / Yüksek", "Clock sync + freshness mask", "Robotics"],
        ["LLM hallucination", "Orta / Yüksek", "RAG + kaynak + komut izolasyonu", "LLM/Safety"],
        ["Dataset leakage", "Orta / Yüksek", "Mission/time split + kör test", "ML"],
        ["ROS/gateway güvenliği", "Orta / Kritik", "Allowlist + signing + network isolation", "Platform"],
    ], [2340, 1780, 3540, 1700])

    doc.add_heading("13. Doğrulama ve kabul kriterleri", level=1)
    doc.add_heading("13.1 Yazılım kabulü", level=2)
    bullet(doc, "Temiz ortamda `python -m unittest discover -s tests -v` başarılı.")
    bullet(doc, "`rovdt demo` veri, model, metrik, karar ve LLM JSONL artefaktlarını üretir.")
    bullet(doc, "Model yükleme sonrası aynı örnek için deterministik sonuç verir.")
    bullet(doc, "Non-nominal veya düşük güvenli durumda otonom yürütme false olur.")
    doc.add_heading("13.2 SIL/HIL kabulü", level=2)
    bullet(doc, "Tüm fault injection senaryolarında beklenen safe state'e geçiş ölçülür.")
    bullet(doc, "End-to-end gecikme ve jitter duty kontrol bütçesinde kalır.")
    bullet(doc, "Sensör dropout/çelişki testleri ve acil surface zinciri başarılıdır.")
    bullet(doc, "Model confidence kalibrasyonu ve sınıf recall eşikleri imzalı test raporunda yer alır.")
    doc.add_heading("13.3 Saha kabulü", level=2)
    bullet(doc, "Önce shadow mode; daha sonra yalnızca sınırlı degraded actions.")
    bullet(doc, "Her wet test için go/no-go checklist, emniyet görevlisi ve geri alma planı.")
    bullet(doc, "Operator override, near-miss ve false alarm'lar yeniden eğitim kuyruğuna girer.")

    page_break(doc)
    doc.add_heading("14. Repo teslimat haritası", level=1)
    add_table(doc, ["Yol", "İçerik"], [
        ["src/rov_dt", "Dataset, schema, classifier, training, advisor, decision, CLI"],
        ["data + artifacts", "1.000 satır örnek telemetry, LLM JSONL, model, metrik ve karar"],
        ["unity/Assets/.../Scripts", "Duty, 6-DOF hydrodynamics ve ML-Agents controller"],
        ["ros2/rov_dt_bridge", "ROS 2 package ve diagnostic node"],
        ["configs + knowledge", "PPO/LoRA ayarları ve operasyon bilgi tabanı"],
        ["tests + .github", "Pipeline testleri ve CI workflow"],
    ], [3200, 6160])
    callout(doc, "İlk çalıştırma", "PYTHONPATH=src ayarlanarak `python -m rov_dt.cli demo --output-dir artifacts/demo` çalıştırılır. Ayrıntılı komutlar README.md içindedir.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
